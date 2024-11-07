from docx import Document
from docx.shared import Inches
import pyttsx3
def speak(text):
    pyttsx3.speak(text)
document =Document()
name =input('what is your name : ')
##################################
speak('hello'+ name + 'how are you today?')
speak('What is your phone number?')
########################################
phone_number =input('what is your phone number : ')
email = input('what is your email : ')
#profile puctur
document.add_picture(
    'fot1.png',
    width=Inches(2.0)
)
#informathin user
document.add_paragraph(
    'name = ' + name + '\n'+'phone number = ' +phone_number + '\n' +'email = ' + email)
#about me
document.add_heading('About me')
about_me = input('Tell about yoursels ? ')
document.add_paragraph( about_me)
#work experience
document.add_heading('work experience')
p=document.add_paragraph()
company = input('enter your company :  ')
from_date = input(' from date : ')
to_date = input('to date : ')
p.add_run(company + ' ').bold = True
p.add_run(from_date + '_' + to_date +'\n').italic =True
experience_details = input('Decibe your experience at '+company +' :')
p.add_run(experience_details)

#mote experiences
while True :
    has_more_experiences = input(
        'do you have more experiences ? yes or no : ').lower()
    if has_more_experiences =='yes':
        p = document.add_paragraph()
        company = input('enter your company :  ')
        from_date = input(' from date : ')
        to_date = input('to date : ')
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '_' + to_date + '\n').italic = True
        experience_details = input('Decibe your experience at ' + company + ' :')
        p.add_run(experience_details)
    elif has_more_experiences == 'no':
        break
    else:
        print('pleas enter yes or no ')
#skills
document.add_heading('SKILLS')
skills = input(' enter skills : ')
p = document.add_paragraph(skills)
p.style ='List Bullet'
while True:
    has_more_skills =input('do you have more skills? yes or no : ').lower()
    if has_more_skills =='yes':
        skills = input(' enter skills : ')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'
    elif has_more_skills =='no':
        break
    else:
        print('pleas enter yes or no ')
#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using marwa kiot kin '
document.save('test.docx')